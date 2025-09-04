"""
多模态AI支持服务
提供图像处理、文档解析等多媒体内容处理能力
"""

import os
import io
import base64
import logging
import mimetypes
from typing import Dict, List, Optional, Any, Union, Tuple
from dataclasses import dataclass
from enum import Enum
from pathlib import Path
import json
import hashlib

# 图像处理相关
try:
    from PIL import Image, ImageDraw, ImageFont
    import cv2
    import numpy as np
    IMAGE_PROCESSING_AVAILABLE = True
except ImportError:
    IMAGE_PROCESSING_AVAILABLE = False
    logging.warning("图像处理库未安装，图像处理功能将不可用")

# 文档处理相关
try:
    import PyPDF2
    import docx
    from docx import Document
    DOCUMENT_PROCESSING_AVAILABLE = True
except ImportError:
    DOCUMENT_PROCESSING_AVAILABLE = False
    logging.warning("文档处理库未安装，文档解析功能将不可用")

from .ai_service import AIService


class MediaType(Enum):
    """媒体类型枚举"""
    IMAGE = "image"
    PDF = "pdf"
    WORD = "word"
    TEXT = "text"
    AUDIO = "audio"
    VIDEO = "video"


class ProcessingType(Enum):
    """处理类型枚举"""
    RECOGNITION = "recognition"      # 识别
    ANALYSIS = "analysis"            # 分析
    EXTRACTION = "extraction"        # 提取
    CONVERSION = "conversion"        # 转换
    ENHANCEMENT = "enhancement"      # 增强


@dataclass
class MediaFile:
    """媒体文件信息"""
    file_path: str
    file_name: str
    file_size: int
    media_type: MediaType
    mime_type: str
    hash_value: str
    metadata: Dict[str, Any] = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


@dataclass
class ProcessingResult:
    """处理结果"""
    success: bool
    media_type: MediaType
    processing_type: ProcessingType
    result_data: Any
    extracted_text: Optional[str] = None
    analysis_result: Optional[Dict[str, Any]] = None
    error_message: Optional[str] = None
    processing_time: float = 0.0


class MultimodalAIService:
    """多模态AI服务"""
    
    def __init__(self, config: Dict[str, Any], ai_service: AIService):
        self.config = config
        self.ai_service = ai_service
        self.logger = logging.getLogger(__name__)
        
        # 检查依赖库
        self._check_dependencies()
        
        # 支持的媒体类型
        self.supported_media_types = config.get('supported_media_types', [
            'image', 'pdf', 'word', 'text'
        ])
        
        # 处理配置
        self.image_config = config.get('image_processing', {})
        self.document_config = config.get('document_processing', {})
        
        # 输出目录
        self.output_dir = config.get('output_directory', './outputs/multimodal')
        os.makedirs(self.output_dir, exist_ok=True)
        
        self.logger.info("多模态AI服务初始化完成")
    
    def _check_dependencies(self):
        """检查依赖库"""
        if not IMAGE_PROCESSING_AVAILABLE:
            self.logger.warning("图像处理功能不可用，请安装: pip install pillow opencv-python numpy")
        
        if not DOCUMENT_PROCESSING_AVAILABLE:
            self.logger.warning("文档处理功能不可用，请安装: pip install PyPDF2 python-docx")
    
    def process_media(self, file_path: str, processing_type: ProcessingType, 
                     options: Dict[str, Any] = None) -> ProcessingResult:
        """
        处理媒体文件
        
        Args:
            file_path: 文件路径
            processing_type: 处理类型
            options: 处理选项
            
        Returns:
            处理结果
        """
        if not os.path.exists(file_path):
            return ProcessingResult(
                success=False,
                media_type=MediaType.TEXT,
                processing_type=processing_type,
                result_data=None,
                error_message=f"文件不存在: {file_path}"
            )
        
        # 识别媒体类型
        media_type = self._detect_media_type(file_path)
        
        if media_type not in self.supported_media_types:
            return ProcessingResult(
                success=False,
                media_type=media_type,
                processing_type=processing_type,
                result_data=None,
                error_message=f"不支持的媒体类型: {media_type}"
            )
        
        # 根据媒体类型和处理类型选择处理方法
        if media_type == MediaType.IMAGE:
            return self._process_image(file_path, processing_type, options)
        elif media_type == MediaType.PDF:
            return self._process_pdf(file_path, processing_type, options)
        elif media_type == MediaType.WORD:
            return self._process_word(file_path, processing_type, options)
        elif media_type == MediaType.TEXT:
            return self._process_text(file_path, processing_type, options)
        else:
            return ProcessingResult(
                success=False,
                media_type=media_type,
                processing_type=processing_type,
                result_data=None,
                error_message=f"未实现的媒体类型处理: {media_type}"
            )
    
    def _detect_media_type(self, file_path: str) -> MediaType:
        """检测媒体类型"""
        mime_type, _ = mimetypes.guess_type(file_path)
        
        if mime_type:
            if mime_type.startswith('image/'):
                return MediaType.IMAGE
            elif mime_type == 'application/pdf':
                return MediaType.PDF
            elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                              'application/msword']:
                return MediaType.WORD
            elif mime_type.startswith('text/'):
                return MediaType.TEXT
        
        # 根据文件扩展名判断
        ext = Path(file_path).suffix.lower()
        if ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp']:
            return MediaType.IMAGE
        elif ext == '.pdf':
            return MediaType.PDF
        elif ext in ['.docx', '.doc']:
            return MediaType.WORD
        elif ext in ['.txt', '.md', '.json', '.xml', '.csv']:
            return MediaType.TEXT
        
        return MediaType.TEXT
    
    def _process_image(self, file_path: str, processing_type: ProcessingType, 
                      options: Dict[str, Any] = None) -> ProcessingResult:
        """处理图像"""
        if not IMAGE_PROCESSING_AVAILABLE:
            return ProcessingResult(
                success=False,
                media_type=MediaType.IMAGE,
                processing_type=processing_type,
                result_data=None,
                error_message="图像处理库未安装"
            )
        
        try:
            import time
            start_time = time.time()
            
            # 加载图像
            image = Image.open(file_path)
            
            if processing_type == ProcessingType.RECOGNITION:
                result = self._recognize_image(image, options)
            elif processing_type == ProcessingType.ANALYSIS:
                result = self._analyze_image(image, options)
            elif processing_type == ProcessingType.EXTRACTION:
                result = self._extract_from_image(image, options)
            elif processing_type == ProcessingType.ENHANCEMENT:
                result = self._enhance_image(image, options)
            else:
                result = self._analyze_image(image, options)  # 默认分析
            
            processing_time = time.time() - start_time
            
            return ProcessingResult(
                success=True,
                media_type=MediaType.IMAGE,
                processing_type=processing_type,
                result_data=result,
                processing_time=processing_time
            )
            
        except Exception as e:
            self.logger.error(f"图像处理失败: {e}")
            return ProcessingResult(
                success=False,
                media_type=MediaType.IMAGE,
                processing_type=processing_type,
                result_data=None,
                error_message=str(e)
            )
    
    def _recognize_image(self, image: Image.Image, options: Dict[str, Any] = None) -> Dict[str, Any]:
        """图像识别"""
        # 将图像转换为base64
        buffer = io.BytesIO()
        image.save(buffer, format='PNG')
        image_base64 = base64.b64encode(buffer.getvalue()).decode()
        
        # 构建提示词
        prompt = options.get('prompt', '请识别这张图片中的内容，包括文字、物体、场景等，并详细描述。')
        
        # 调用AI服务进行识别
        try:
            response = self.ai_service.generate_text_with_image(
                prompt=prompt,
                image_base64=image_base64,
                max_tokens=options.get('max_tokens', 1000),
                temperature=options.get('temperature', 0.3)
            )
            
            return {
                'recognition_result': response,
                'image_size': image.size,
                'image_mode': image.mode,
                'processing_options': options
            }
            
        except Exception as e:
            self.logger.error(f"AI图像识别失败: {e}")
            return {
                'recognition_result': f"图像识别失败: {str(e)}",
                'image_size': image.size,
                'image_mode': image.mode,
                'error': str(e)
            }
    
    def _analyze_image(self, image: Image.Image, options: Dict[str, Any] = None) -> Dict[str, Any]:
        """图像分析"""
        # 基本图像信息
        analysis_result = {
            'image_size': image.size,
            'image_mode': image.mode,
            'file_format': image.format,
            'color_palette': self._analyze_color_palette(image),
            'brightness': self._analyze_brightness(image),
            'contrast': self._analyze_contrast(image)
        }
        
        # 使用OpenCV进行更详细的分析
        if hasattr(cv2, 'imread'):
            try:
                cv_image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
                analysis_result.update(self._cv2_analyze_image(cv_image))
            except Exception as e:
                self.logger.warning(f"OpenCV分析失败: {e}")
        
        # AI分析
        if options.get('use_ai_analysis', True):
            ai_analysis = self._recognize_image(image, options)
            analysis_result['ai_analysis'] = ai_analysis.get('recognition_result', '')
        
        return analysis_result
    
    def _analyze_color_palette(self, image: Image.Image) -> Dict[str, Any]:
        """分析颜色调色板"""
        # 转换为RGB模式
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # 获取主要颜色
        colors = image.getcolors(maxcolors=256)
        if colors:
            colors.sort(key=lambda x: x[0], reverse=True)
            top_colors = colors[:5]
            
            return {
                'top_colors': [(f"#{c[1][0]:02x}{c[1][1]:02x}{c[1][2]:02x}", c[0]) for c in top_colors],
                'total_colors': len(colors)
            }
        
        return {'top_colors': [], 'total_colors': 0}
    
    def _analyze_brightness(self, image: Image.Image) -> float:
        """分析亮度"""
        if image.mode != 'L':
            image = image.convert('L')
        
        pixels = list(image.getdata())
        return sum(pixels) / len(pixels) / 255.0
    
    def _analyze_contrast(self, image: Image.Image) -> float:
        """分析对比度"""
        if image.mode != 'L':
            image = image.convert('L')
        
        pixels = list(image.getdata())
        mean = sum(pixels) / len(pixels)
        variance = sum((p - mean) ** 2 for p in pixels) / len(pixels)
        return (variance ** 0.5) / 255.0
    
    def _cv2_analyze_image(self, cv_image) -> Dict[str, Any]:
        """使用OpenCV分析图像"""
        analysis = {}
        
        # 边缘检测
        gray = cv2.cvtColor(cv_image, cv2.COLOR_BGR2GRAY)
        edges = cv2.Canny(gray, 50, 150)
        analysis['edge_density'] = np.sum(edges > 0) / (edges.shape[0] * edges.shape[1])
        
        # 直方图分析
        hist = cv2.calcHist([cv_image], [0, 1, 2], None, [256, 256, 256], [0, 256, 0, 256, 0, 256])
        analysis['histogram_entropy'] = -np.sum(hist * np.log2(hist + 1e-10))
        
        return analysis
    
    def _extract_from_image(self, image: Image.Image, options: Dict[str, Any] = None) -> Dict[str, Any]:
        """从图像中提取信息"""
        extraction_result = {
            'text_extraction': self._extract_text_from_image(image, options),
            'metadata_extraction': self._extract_metadata_from_image(image),
            'feature_extraction': self._extract_features_from_image(image, options)
        }
        
        return extraction_result
    
    def _extract_text_from_image(self, image: Image.Image, options: Dict[str, Any] = None) -> str:
        """从图像中提取文字"""
        prompt = options.get('text_extraction_prompt', 
                           '请提取这张图片中的所有文字内容，包括标题、正文、标签等，按顺序输出。')
        
        # 使用AI服务提取文字
        try:
            buffer = io.BytesIO()
            image.save(buffer, format='PNG')
            image_base64 = base64.b64encode(buffer.getvalue()).decode()
            
            response = self.ai_service.generate_text_with_image(
                prompt=prompt,
                image_base64=image_base64,
                max_tokens=options.get('max_tokens', 500),
                temperature=options.get('temperature', 0.1)
            )
            
            return response
            
        except Exception as e:
            self.logger.error(f"文字提取失败: {e}")
            return f"文字提取失败: {str(e)}"
    
    def _extract_metadata_from_image(self, image: Image.Image) -> Dict[str, Any]:
        """提取图像元数据"""
        metadata = {
            'format': image.format,
            'mode': image.mode,
            'size': image.size,
            'width': image.width,
            'height': image.height
        }
        
        # 提取EXIF信息
        if hasattr(image, '_getexif') and image._getexif():
            metadata['exif'] = dict(image._getexif())
        
        return metadata
    
    def _extract_features_from_image(self, image: Image.Image, options: Dict[str, Any] = None) -> Dict[str, Any]:
        """提取图像特征"""
        features = {
            'basic_features': {
                'aspect_ratio': image.width / image.height,
                'pixel_count': image.width * image.height,
                'file_size_estimate': image.width * image.height * 3  # 估算文件大小
            }
        }
        
        # 使用OpenCV提取更多特征
        if hasattr(cv2, 'imread'):
            try:
                cv_image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
                
                # 颜色特征
                hsv = cv2.cvtColor(cv_image, cv2.COLOR_BGR2HSV)
                features['color_features'] = {
                    'hue_mean': np.mean(hsv[:, :, 0]),
                    'saturation_mean': np.mean(hsv[:, :, 1]),
                    'value_mean': np.mean(hsv[:, :, 2])
                }
                
                # 纹理特征
                gray = cv2.cvtColor(cv_image, cv2.COLOR_BGR2GRAY)
                features['texture_features'] = {
                    'std_dev': np.std(gray),
                    'variance': np.var(gray)
                }
                
            except Exception as e:
                self.logger.warning(f"特征提取失败: {e}")
        
        return features
    
    def _enhance_image(self, image: Image.Image, options: Dict[str, Any] = None) -> Dict[str, Any]:
        """图像增强"""
        enhanced_image = image.copy()
        
        # 应用增强选项
        if options.get('enhance_brightness'):
            enhanced_image = self._enhance_brightness(enhanced_image, options.get('brightness_factor', 1.2))
        
        if options.get('enhance_contrast'):
            enhanced_image = self._enhance_contrast(enhanced_image, options.get('contrast_factor', 1.5))
        
        if options.get('enhance_sharpness'):
            enhanced_image = self._enhance_sharpness(enhanced_image)
        
        # 保存增强后的图像
        output_path = os.path.join(self.output_dir, f"enhanced_{Path(image.filename).name}")
        enhanced_image.save(output_path)
        
        return {
            'original_size': image.size,
            'enhanced_size': enhanced_image.size,
            'output_path': output_path,
            'enhancement_options': options
        }
    
    def _enhance_brightness(self, image: Image.Image, factor: float) -> Image.Image:
        """增强亮度"""
        from PIL import ImageEnhance
        enhancer = ImageEnhance.Brightness(image)
        return enhancer.enhance(factor)
    
    def _enhance_contrast(self, image: Image.Image, factor: float) -> Image.Image:
        """增强对比度"""
        from PIL import ImageEnhance
        enhancer = ImageEnhance.Contrast(image)
        return enhancer.enhance(factor)
    
    def _enhance_sharpness(self, image: Image.Image) -> Image.Image:
        """增强锐度"""
        from PIL import ImageEnhance
        enhancer = ImageEnhance.Sharpness(image)
        return enhancer.enhance(2.0)
    
    def _process_pdf(self, file_path: str, processing_type: ProcessingType, 
                    options: Dict[str, Any] = None) -> ProcessingResult:
        """处理PDF文档"""
        if not DOCUMENT_PROCESSING_AVAILABLE:
            return ProcessingResult(
                success=False,
                media_type=MediaType.PDF,
                processing_type=processing_type,
                result_data=None,
                error_message="PDF处理库未安装"
            )
        
        try:
            import time
            start_time = time.time()
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                if processing_type == ProcessingType.EXTRACTION:
                    result = self._extract_from_pdf(pdf_reader, options)
                elif processing_type == ProcessingType.ANALYSIS:
                    result = self._analyze_pdf(pdf_reader, options)
                elif processing_type == ProcessingType.CONVERSION:
                    result = self._convert_pdf(pdf_reader, options)
                else:
                    result = self._extract_from_pdf(pdf_reader, options)  # 默认提取
            
            processing_time = time.time() - start_time
            
            return ProcessingResult(
                success=True,
                media_type=MediaType.PDF,
                processing_type=processing_type,
                result_data=result,
                extracted_text=result.get('text_content', ''),
                processing_time=processing_time
            )
            
        except Exception as e:
            self.logger.error(f"PDF处理失败: {e}")
            return ProcessingResult(
                success=False,
                media_type=MediaType.PDF,
                processing_type=processing_type,
                result_data=None,
                error_message=str(e)
            )
    
    def _extract_from_pdf(self, pdf_reader, options: Dict[str, Any] = None) -> Dict[str, Any]:
        """从PDF中提取内容"""
        text_content = ""
        metadata = {}
        
        # 提取文本
        for page_num, page in enumerate(pdf_reader.pages):
            text_content += f"\n--- 第 {page_num + 1} 页 ---\n"
            text_content += page.extract_text()
        
        # 提取元数据
        if pdf_reader.metadata:
            metadata = dict(pdf_reader.metadata)
        
        # 使用AI分析提取的内容
        if options.get('use_ai_analysis', True) and text_content.strip():
            ai_analysis = self._analyze_text_with_ai(text_content, options)
        else:
            ai_analysis = {}
        
        return {
            'text_content': text_content,
            'metadata': metadata,
            'page_count': len(pdf_reader.pages),
            'ai_analysis': ai_analysis
        }
    
    def _analyze_pdf(self, pdf_reader, options: Dict[str, Any] = None) -> Dict[str, Any]:
        """分析PDF文档"""
        analysis = {
            'page_count': len(pdf_reader.pages),
            'metadata': dict(pdf_reader.metadata) if pdf_reader.metadata else {},
            'file_size': os.path.getsize(pdf_reader.stream.name) if hasattr(pdf_reader.stream, 'name') else 0
        }
        
        # 提取文本进行分析
        text_content = ""
        for page in pdf_reader.pages:
            text_content += page.extract_text()
        
        if text_content.strip():
            analysis['text_analysis'] = self._analyze_text_content(text_content)
            analysis['ai_analysis'] = self._analyze_text_with_ai(text_content, options)
        
        return analysis
    
    def _convert_pdf(self, pdf_reader, options: Dict[str, Any] = None) -> Dict[str, Any]:
        """转换PDF文档"""
        # 提取文本内容
        text_content = ""
        for page in pdf_reader.pages:
            text_content += page.extract_text()
        
        # 保存为文本文件
        output_path = os.path.join(self.output_dir, f"converted_{Path(pdf_reader.stream.name).stem}.txt")
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text_content)
        
        return {
            'output_path': output_path,
            'text_content': text_content,
            'conversion_format': 'text'
        }
    
    def _process_word(self, file_path: str, processing_type: ProcessingType, 
                     options: Dict[str, Any] = None) -> ProcessingResult:
        """处理Word文档"""
        if not DOCUMENT_PROCESSING_AVAILABLE:
            return ProcessingResult(
                success=False,
                media_type=MediaType.WORD,
                processing_type=processing_type,
                result_data=None,
                error_message="Word文档处理库未安装"
            )
        
        try:
            import time
            start_time = time.time()
            
            doc = Document(file_path)
            
            if processing_type == ProcessingType.EXTRACTION:
                result = self._extract_from_word(doc, options)
            elif processing_type == ProcessingType.ANALYSIS:
                result = self._analyze_word(doc, options)
            elif processing_type == ProcessingType.CONVERSION:
                result = self._convert_word(doc, options)
            else:
                result = self._extract_from_word(doc, options)  # 默认提取
            
            processing_time = time.time() - start_time
            
            return ProcessingResult(
                success=True,
                media_type=MediaType.WORD,
                processing_type=processing_type,
                result_data=result,
                extracted_text=result.get('text_content', ''),
                processing_time=processing_time
            )
            
        except Exception as e:
            self.logger.error(f"Word文档处理失败: {e}")
            return ProcessingResult(
                success=False,
                media_type=MediaType.WORD,
                processing_type=processing_type,
                result_data=None,
                error_message=str(e)
            )
    
    def _extract_from_word(self, doc, options: Dict[str, Any] = None) -> Dict[str, Any]:
        """从Word文档中提取内容"""
        text_content = ""
        paragraphs = []
        
        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                text_content += para.text + "\n"
                paragraphs.append(para.text)
        
        # 提取表格
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        # 使用AI分析
        if options.get('use_ai_analysis', True) and text_content.strip():
            ai_analysis = self._analyze_text_with_ai(text_content, options)
        else:
            ai_analysis = {}
        
        return {
            'text_content': text_content,
            'paragraphs': paragraphs,
            'tables': tables,
            'paragraph_count': len(paragraphs),
            'table_count': len(tables),
            'ai_analysis': ai_analysis
        }
    
    def _analyze_word(self, doc, options: Dict[str, Any] = None) -> Dict[str, Any]:
        """分析Word文档"""
        # 提取文本内容
        text_content = ""
        for para in doc.paragraphs:
            text_content += para.text + "\n"
        
        analysis = {
            'paragraph_count': len([p for p in doc.paragraphs if p.text.strip()]),
            'table_count': len(doc.tables),
            'section_count': len(doc.sections),
            'text_analysis': self._analyze_text_content(text_content)
        }
        
        if text_content.strip():
            analysis['ai_analysis'] = self._analyze_text_with_ai(text_content, options)
        
        return analysis
    
    def _convert_word(self, doc, options: Dict[str, Any] = None) -> Dict[str, Any]:
        """转换Word文档"""
        # 提取文本内容
        text_content = ""
        for para in doc.paragraphs:
            text_content += para.text + "\n"
        
        # 保存为文本文件
        output_path = os.path.join(self.output_dir, f"converted_{Path(doc._path).stem}.txt")
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text_content)
        
        return {
            'output_path': output_path,
            'text_content': text_content,
            'conversion_format': 'text'
        }
    
    def _process_text(self, file_path: str, processing_type: ProcessingType, 
                     options: Dict[str, Any] = None) -> ProcessingResult:
        """处理文本文件"""
        try:
            import time
            start_time = time.time()
            
            with open(file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            if processing_type == ProcessingType.ANALYSIS:
                result = self._analyze_text_content(text_content)
            elif processing_type == ProcessingType.EXTRACTION:
                result = self._extract_from_text(text_content, options)
            else:
                result = self._analyze_text_content(text_content)  # 默认分析
            
            processing_time = time.time() - start_time
            
            return ProcessingResult(
                success=True,
                media_type=MediaType.TEXT,
                processing_type=processing_type,
                result_data=result,
                extracted_text=text_content,
                processing_time=processing_time
            )
            
        except Exception as e:
            self.logger.error(f"文本处理失败: {e}")
            return ProcessingResult(
                success=False,
                media_type=MediaType.TEXT,
                processing_type=processing_type,
                result_data=None,
                error_message=str(e)
            )
    
    def _analyze_text_content(self, text_content: str) -> Dict[str, Any]:
        """分析文本内容"""
        analysis = {
            'character_count': len(text_content),
            'word_count': len(text_content.split()),
            'line_count': len(text_content.splitlines()),
            'paragraph_count': len([p for p in text_content.split('\n\n') if p.strip()]),
            'average_word_length': sum(len(word) for word in text_content.split()) / max(len(text_content.split()), 1),
            'unique_words': len(set(word.lower() for word in text_content.split() if word.isalpha()))
        }
        
        return analysis
    
    def _extract_from_text(self, text_content: str, options: Dict[str, Any] = None) -> Dict[str, Any]:
        """从文本中提取信息"""
        extraction = {
            'text_content': text_content,
            'text_analysis': self._analyze_text_content(text_content)
        }
        
        # 使用AI提取关键信息
        if options.get('use_ai_extraction', True):
            ai_extraction = self._extract_with_ai(text_content, options)
            extraction['ai_extraction'] = ai_extraction
        
        return extraction
    
    def _analyze_text_with_ai(self, text_content: str, options: Dict[str, Any] = None) -> Dict[str, Any]:
        """使用AI分析文本"""
        prompt = options.get('analysis_prompt', 
                           '请分析以下文本内容，包括主题、关键信息、情感倾向、写作风格等：\n\n{text}')
        
        try:
            response = self.ai_service.generate_text(
                prompt=prompt.format(text=text_content[:2000]),  # 限制长度
                max_tokens=options.get('max_tokens', 500),
                temperature=options.get('temperature', 0.3)
            )
            
            return {
                'ai_analysis': response,
                'analysis_type': 'comprehensive'
            }
            
        except Exception as e:
            self.logger.error(f"AI文本分析失败: {e}")
            return {
                'ai_analysis': f"AI分析失败: {str(e)}",
                'analysis_type': 'error'
            }
    
    def _extract_with_ai(self, text_content: str, options: Dict[str, Any] = None) -> Dict[str, Any]:
        """使用AI提取信息"""
        prompt = options.get('extraction_prompt', 
                           '请从以下文本中提取关键信息，包括主要观点、重要数据、关键人物等：\n\n{text}')
        
        try:
            response = self.ai_service.generate_text(
                prompt=prompt.format(text=text_content[:2000]),
                max_tokens=options.get('max_tokens', 300),
                temperature=options.get('temperature', 0.2)
            )
            
            return {
                'extracted_info': response,
                'extraction_type': 'key_information'
            }
            
        except Exception as e:
            self.logger.error(f"AI信息提取失败: {e}")
            return {
                'extracted_info': f"信息提取失败: {str(e)}",
                'extraction_type': 'error'
            }
    
    def batch_process(self, file_paths: List[str], processing_type: ProcessingType, 
                     options: Dict[str, Any] = None) -> List[ProcessingResult]:
        """批量处理文件"""
        results = []
        
        for file_path in file_paths:
            result = self.process_media(file_path, processing_type, options)
            results.append(result)
        
        return results
    
    def get_supported_formats(self) -> Dict[str, List[str]]:
        """获取支持的格式"""
        return {
            'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp'],
            'pdf': ['.pdf'],
            'word': ['.docx', '.doc'],
            'text': ['.txt', '.md', '.json', '.xml', '.csv']
        }
    
    def get_processing_capabilities(self) -> Dict[str, List[str]]:
        """获取处理能力"""
        capabilities = {
            'image': ['recognition', 'analysis', 'extraction', 'enhancement'],
            'pdf': ['extraction', 'analysis', 'conversion'],
            'word': ['extraction', 'analysis', 'conversion'],
            'text': ['analysis', 'extraction']
        }
        
        # 根据可用库调整能力
        if not IMAGE_PROCESSING_AVAILABLE:
            capabilities['image'] = []
        
        if not DOCUMENT_PROCESSING_AVAILABLE:
            capabilities['pdf'] = []
            capabilities['word'] = []
        
        return capabilities
